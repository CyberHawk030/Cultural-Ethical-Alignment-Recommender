import os
import asyncio
import re
import json
from typing import List, Optional
from datetime import datetime

# --- FastAPI Imports ---
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, ValidationError

# --- PDF/DOCX Parsing Imports ---
from PyPDF2 import PdfReader
import docx

# --- AI & LangChain Imports ---
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.tools import Tool
from crewai import Agent, Task, Crew, Process
from langchain.agents import load_tools

# --- MongoDB Imports ---
from bson import ObjectId
from database import report_collection
from models import ReportSchema, ReportSummary

# --------------------------------------------------------------------------
# --- 1. FASTAPI APP INITIALIZATION & CORS CONFIGURATION ---
# --------------------------------------------------------------------------

app = FastAPI(
    title="Ethical Compass API",
    description="API for running the multi-agent ethical alignment system.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --------------------------------------------------------------------------
# --- 2. IN-MEMORY STORAGE & HELPER FUNCTIONS ---
# --------------------------------------------------------------------------

knowledge_vector_store = None

class ApiKeyRequest(BaseModel):
    google_api_key: str

def parse_files(files: List[UploadFile]) -> str:
    raw_text = ""
    for file in files:
        file.file.seek(0)
        if file.content_type == "application/pdf":
            pdf_reader = PdfReader(file.file)
            for page in pdf_reader.pages:
                raw_text += page.extract_text() or ""
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file.file)
            raw_text += "\n".join([para.text for para in doc.paragraphs])
        elif file.content_type == "text/plain":
            raw_text += file.file.read().decode("utf-8")
    return raw_text

def create_vector_store(text_chunks: List[str]):
    global knowledge_vector_store
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        documents = [Document(page_content=chunk) for chunk in text_chunks]
        knowledge_vector_store = FAISS.from_documents(documents, embedding=embeddings)
        return True
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create vector store: {str(e)}. Is the API Key correct?")

# --------------------------------------------------------------------------
# --- 3. API ENDPOINTS ---
# --------------------------------------------------------------------------

@app.post("/configure")
async def configure_api_key(request: ApiKeyRequest):
    if not request.google_api_key:
        raise HTTPException(status_code=400, detail="API key not provided.")
    os.environ["GOOGLE_API_KEY"] = request.google_api_key
    return {"message": "API Key configured successfully."}

@app.post("/upload")
async def upload_knowledge_base(files: List[UploadFile] = File(...)):
    raw_text = parse_files(files)
    if not raw_text:
        raise HTTPException(status_code=400, detail="No text could be extracted from the uploaded files.")
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_text(raw_text)
    
    create_vector_store(chunks)
    
    return {"message": "Knowledge base processed and ready.", "file_names": [f.filename for f in files]}

@app.post("/analyze")
async def analyze_candidate(
    candidate_name: str = Form(...),
    candidate_written_submissions: str = Form(""),
    critical_penalty: int = Form(...),
    minor_penalty: int = Form(...),
    candidate_files: Optional[List[UploadFile]] = File(None)
):
    if not knowledge_vector_store:
        raise HTTPException(status_code=400, detail="Knowledge base is not initialized.")

    def knowledge_base_search(query: str) -> str:
        return "\n---\n".join([doc.page_content for doc in knowledge_vector_store.similarity_search(query, k=3)])

    knowledge_base_tool = Tool(name="Knowledge Base Search", func=knowledge_base_search, description="Searches the unified knowledge base.")
    
    uploaded_text = parse_files(candidate_files) if candidate_files else ""
    candidate_full_text = candidate_written_submissions + "\n" + uploaded_text
    
    if not candidate_full_text.strip():
        raise HTTPException(status_code=400, detail="No candidate text provided.")

    llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash-latest", temperature=0.1)
    reporting_tools = load_tools(["llm-math"], llm=llm)

    profiler_agent = Agent(role="Profiler Agent", goal="Create structured JSON profiles for the institution and the candidate.", backstory="Expert analyst, distills complex documents into clear, machine-readable data.", llm=llm, tools=[knowledge_base_tool], verbose=True, allow_delegation=False)
    conflict_detector = Agent(role="Conflict Detector Agent", goal="Compare JSON profiles to find misalignments.", backstory="Meticulous logic engine.", llm=llm, verbose=True, allow_delegation=False)
    scenario_evaluator = Agent(role="Case Scenario Evaluator Agent", goal="Generate and evaluate ethical dilemmas.", backstory="Expert in ethical pedagogy.", llm=llm, verbose=True, allow_delegation=False)
    report_generator = Agent(role="Final Report Generator Agent", goal="Synthesize analyses into a comprehensive report.", backstory="Expert HR strategist who uses a calculator for all math.", llm=llm, tools=reporting_tools, verbose=True, allow_delegation=False)
    
    task_profiling = Task(description=f"First, use 'Knowledge Base Search' to understand institutional values. Second, analyze '{candidate_name}'s submissions: \n\n{candidate_full_text}\n\nCreate two separate JSON objects for the institution and candidate profiles.", expected_output="A single JSON object with two keys: 'institution_profile' and 'candidate_profile'.", agent=profiler_agent)
    task_conflict_detection = Task(description="Compare the institution's JSON profile with the candidate's. Identify value clashes.", expected_output="JSON with a 'conflicts' key, listing objects with 'area', 'severity', and 'explanation'.", context=[task_profiling], agent=conflict_detector)
    task_scenario_evaluation = Task(description="From the conflict report, take the most 'Critical' conflict and generate THREE distinct case scenarios.", expected_output="JSON with a 'scenarios' key, a list of 3 objects with 'scenario_title', 'dilemma', 'predicted_response', 'alignment_score' (1-10), 'reasoning_strengths', and 'reasoning_weaknesses'.", context=[task_conflict_detection], agent=scenario_evaluator)
    
    task_final_report = Task(description=f"Synthesize all outputs into a final report for '{candidate_name}'. Use the 'Calculator' to compute the final score: 100 - (critical_penalty * num_critical) - (minor_penalty * num_minor).", expected_output="A JSON object with two keys: 'final_score' (number) and 'report_markdown' (string).", context=[task_profiling, task_conflict_detection, task_scenario_evaluation], agent=report_generator)

    ethical_crew = Crew(agents=[profiler_agent, conflict_detector, scenario_evaluator, report_generator], tasks=[task_profiling, task_conflict_detection, task_scenario_evaluation, task_final_report], process=Process.sequential, verbose=2)
    
    try:
        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(None, ethical_crew.kickoff)
        
        if not result:
             raise HTTPException(status_code=500, detail="Analysis failed: agent timed out.")

        try:
            match = re.search(r'\{.*\}', result, re.DOTALL)
            if not match:
                raise ValueError("No JSON object found in agent output.")
            json_string = match.group(0)
            report_json = json.loads(json_string)
            report_score = int(report_json['final_score'])
            full_report_text = report_json['report_markdown']
        except (json.JSONDecodeError, KeyError, ValueError) as e:
            print(f"Error parsing agent's JSON output: {e}\n--- Full Agent Output ---\n{result}\n-------------------------")
            raise HTTPException(status_code=500, detail="Failed to parse the analysis report from the agent.")

        report_data = ReportSchema(candidate_name=candidate_name, score=report_score, date=datetime.utcnow(), full_report=full_report_text)
        await report_collection.insert_one(report_data.model_dump(by_alias=True, exclude=["id"]))
        
        return {"report": full_report_text}
        
    except Exception as e:
        print(f"An error occurred during agent kickoff: {e}")
        raise HTTPException(status_code=500, detail=f"An error occurred during analysis: {str(e)}")

# ROBUST VERSION of the get_history endpoint
@app.get("/history")
async def get_history():
    try:

        cursor = report_collection.find({}, {"candidate_name": 1, "score": 1, "date": 1, "_id": 1}).sort("date", -1)
        history = []
        async for doc in cursor:
            try:
                history.append({
                    "id": str(doc["_id"]),
                    "candidate_name": doc.get("candidate_name", "Unknown"),
                    "score": doc.get("score", 0),
                    "date": doc.get("date", datetime.utcnow()).isoformat()
                })
            except Exception as e:
                # If a specific document has an issue (e.g., missing _id), skip it
                print(f"Skipping corrupted document in history. Reason: {e}")
                continue
        return history
    except Exception as e:
        print(f"An unexpected error occurred in get_history: {e}")
        raise HTTPException(status_code=500, detail="An unexpected server error occurred while fetching history.")


# ROBUST VERSION of the get_report endpoint to prevent Pydantic crash
@app.get("/report/{report_id}")
async def get_report(report_id: str):
    try:
        if not ObjectId.is_valid(report_id):
            raise HTTPException(status_code=400, detail="Invalid report ID format.")
        
        report_doc = await report_collection.find_one({"_id": ObjectId(report_id)})
        
        if report_doc is None:
            raise HTTPException(status_code=404, detail="Report not found.")
        
        # ** THE FIX IS HERE: Manually build the response dictionary **
        # This bypasses the Pydantic v2 validation bug by not using model_validate
        response_data = {
            "id": str(report_doc["_id"]),
            "candidate_name": report_doc.get("candidate_name", "Unknown"),
            "score": report_doc.get("score", 0),
            "date": report_doc.get("date", datetime.utcnow()).isoformat(),
            "full_report": report_doc.get("full_report", "Report content is not available for this legacy entry.")
        }
        
        # Return a JSONResponse directly. This gives us full control and avoids auto-validation.
        return JSONResponse(content=response_data)

    except Exception as e:
        print(f"An unexpected error occurred in get_report for ID {report_id}: {e}")
        raise HTTPException(status_code=500, detail="An unexpected server error occurred.")
